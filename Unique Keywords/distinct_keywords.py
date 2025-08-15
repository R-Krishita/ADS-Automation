import ads
from docx import Document
import re

# === SETUP ===
# Add your ADS API key
ads.config.token = 'LBkyAfNatpoqljbX24TxS7AHFrAenKWAtGlDf0vL'

# List of ADS abstract URLs
urls = [
    "https://ui.adsabs.harvard.edu/abs/2018SSRv..214...46G/abstract",
    "https://ui.adsabs.harvard.edu/abs/2002SoPh..208..297V/abstract",
    "https://ui.adsabs.harvard.edu/abs/2011SSRv..159...19F/abstract",
    "https://ui.adsabs.harvard.edu/abs/2010ApJ...719..655V/abstract",
    "https://ui.adsabs.harvard.edu/abs/2019JSWSC...9A...6K/abstract",
    "https://ui.adsabs.harvard.edu/abs/2009A%26A...506.1429C/abstract",
    "https://ui.adsabs.harvard.edu/abs/2017ApJ...844...54D/abstract",
    "https://ui.adsabs.harvard.edu/abs/2020ApJ...896..119K/abstract",
    "https://ui.adsabs.harvard.edu/abs/2015JSWSC...5A..23R/abstract",
    "https://ui.adsabs.harvard.edu/abs/2009SSRv..147..121M/abstract",
    "https://ui.adsabs.harvard.edu/abs/2017SoPh..292..169P/abstract",
    "https://ui.adsabs.harvard.edu/abs/2018JSWSC...8A...9S/abstract",
    "https://ui.adsabs.harvard.edu/abs/2023AdSpR..72.5161W/abstract",
    "https://ui.adsabs.harvard.edu/abs/2020WDMKD..10.1349F/abstract",
    "https://ui.adsabs.harvard.edu/abs/2019ApJ...877...67B/abstract",
    "https://ui.adsabs.harvard.edu/abs/2016ApJ...823...41D/abstract",
    "https://ui.adsabs.harvard.edu/abs/2015SoPh..290.3425J/abstract",
    "https://ui.adsabs.harvard.edu/abs/2017A%26A...601A.125P/abstract", 
    "https://ui.adsabs.harvard.edu/abs/2019ApJ...881...15W/abstract",
    "https://ui.adsabs.harvard.edu/abs/2021LRSP...18....4T/abstract",
    "https://ui.adsabs.harvard.edu/abs/2023NatAs...7.1171J/abstract"
]

# === EXTRACT BIBCODES FROM URLS ===
bibcodes = [re.search(r'abs/([^/]+)', url).group(1) for url in urls]

# === START DOCX ===
doc = Document()
doc.add_heading('ADS Paper Titles and Keywords', 0)

# Table with title and keywords
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Paper Title'
hdr_cells[1].text = 'Keywords'

# Set for storing unique keywords
all_keywords = set()

# === LOOP THROUGH PAPERS ===
for bibcode in bibcodes:
    papers = list(ads.SearchQuery(bibcode=bibcode, fl=['title', 'keyword']))
    if not papers:
        continue
    paper = papers[0]

    title = paper.title[0] if paper.title else 'No title found'
    keywords = paper.keyword if paper.keyword else []

    # Add to table
    row_cells = table.add_row().cells
    row_cells[0].text = title
    row_cells[1].text = ', '.join(keywords)

    # Update keyword set
    all_keywords.update([kw.strip().lower() for kw in keywords])

# === ADD UNIQUE KEYWORDS LIST TO DOCX ===
doc.add_page_break()
doc.add_heading('Distinct Keywords (Sorted)', level=1)

for kw in sorted(all_keywords):
    doc.add_paragraph(kw, style='List Bullet')

# === SAVE DOCX ===
doc.save('Group 7.docx')
print("Saved Group 7.docx successfully.")
