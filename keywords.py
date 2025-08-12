import ads
from docx import Document
import re

# Set your ADS API key
ads.config.token = 'LBkyAfNatpoqljbX24TxS7AHFrAenKWAtGlDf0vL'

# List of ADS abstract URLs
urls = [
    "https://ui.adsabs.harvard.edu/abs/2022ApJ...931..163S/abstract",
    "https://ui.adsabs.harvard.edu/abs/2020ApJ...895....3W/abstract",
    "https://ui.adsabs.harvard.edu/abs/2019SpWea..17.1404C/abstract",
    "https://ui.adsabs.harvard.edu/abs/2020SpWea..1802440J/abstract",
    "https://ui.adsabs.harvard.edu/abs/2022SpWea..2002842K/abstract",
    "https://ui.adsabs.harvard.edu/abs/2019ApJS..242....7G/abstract",
    "https://ui.adsabs.harvard.edu/abs/2019arXiv191212360S/abstract",
    "https://ui.adsabs.harvard.edu/abs/2021JSWSC..11...42G/abstract",
    "https://ui.adsabs.harvard.edu/abs/2021AdSpR..68.1377X/abstract",
    "https://ui.adsabs.harvard.edu/abs/2022ApJ...928..157W/abstract",
    "https://ui.adsabs.harvard.edu/abs/2021A%26A...648A..53D/abstract",
    "https://ui.adsabs.harvard.edu/abs/2023AdSpR..72.5161W/abstract",
    "https://ui.adsabs.harvard.edu/abs/2019JGRA..124..790C/abstract",
    "https://ui.adsabs.harvard.edu/abs/2018SpWea..16.1583L/abstract",
    "https://ui.adsabs.harvard.edu/abs/2019ApJ...884..175W/abstract",
    "https://ui.adsabs.harvard.edu/abs/2019SpWea..17.1166C/abstract",
    "https://ui.adsabs.harvard.edu/abs/2018SoPh..293...48J/abstract",
    "https://ui.adsabs.harvard.edu/abs/2018SSRv..214...46G/abstract",
    "https://ui.adsabs.harvard.edu/abs/2019ApJ...881...15W/abstract",
    "https://ui.adsabs.harvard.edu/abs/2018ApJ...869...13J/abstract",
    "https://ui.adsabs.harvard.edu/abs/2021ApJS..257...50T/abstract"
]

# Extract bibcodes from URLs using regex
bibcodes = [re.search(r'abs/([^/]+)', url).group(1) for url in urls]

# Create a Word document
doc = Document()
doc.add_heading('Group 3', 0)

# Add table
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Paper Title'
hdr_cells[1].text = 'Keywords'

# Fetch paper info using ADS API
for bibcode in bibcodes:
    papers = list(ads.SearchQuery(bibcode=bibcode, fl=['title', 'keyword']))
    if not papers:
        continue
    paper = papers[0]

    title = paper.title[0] if paper.title else 'No title found'
    keywords = ', '.join(paper.keyword) if paper.keyword else 'No keywords found'

    row_cells = table.add_row().cells
    row_cells[0].text = title
    row_cells[1].text = keywords

# Save docx
doc.save('Group 3.docx')
print("Saved Group 3.docx successfully.")
